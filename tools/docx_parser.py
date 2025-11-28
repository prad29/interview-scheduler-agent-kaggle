"""
DOCX Parser

Extracts text content from Microsoft Word DOCX files.
"""

import logging
from pathlib import Path
from typing import Optional, List

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.error("python-docx not available. Install with: pip install python-docx")


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text content from DOCX file
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Extracted text content
        
    Raises:
        ImportError: If python-docx is not installed
        FileNotFoundError: If DOCX file doesn't exist
        Exception: If extraction fails
    """
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is required to parse DOCX files. "
            "Install it with: pip install python-docx"
        )
    
    # Validate file exists
    docx_file = Path(docx_path)
    if not docx_file.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")
    
    logger.info(f"Extracting text from DOCX: {docx_path}")
    
    try:
        doc = Document(docx_path)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            table_text = extract_table_text(table)
            if table_text:
                text_parts.extend(table_text)
        
        # Extract text from headers and footers
        for section in doc.sections:
            # Header
            header = section.header
            for paragraph in header.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Footer
            footer = section.footer
            for paragraph in footer.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
        
        full_text = '\n'.join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from DOCX")
        
        return full_text.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {docx_path}. Error: {str(e)}")


def extract_table_text(table: 'Table') -> List[str]:
    """
    Extract text from a Word table
    
    Args:
        table: DOCX Table object
        
    Returns:
        List of text strings from table rows
    """
    table_text = []
    
    try:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            
            if row_text:
                # Join cells with pipe separator for readability
                table_text.append(' | '.join(row_text))
        
        return table_text
        
    except Exception as e:
        logger.warning(f"Error extracting table text: {str(e)}")
        return []


def get_docx_metadata(docx_path: str) -> dict:
    """
    Extract metadata from DOCX file
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Dictionary with DOCX metadata
    """
    if not DOCX_AVAILABLE:
        return {}
    
    metadata = {
        'author': None,
        'title': None,
        'subject': None,
        'keywords': None,
        'created': None,
        'modified': None,
        'last_modified_by': None,
        'paragraphs': 0,
        'tables': 0
    }
    
    try:
        doc = Document(docx_path)
        
        # Core properties
        core_props = doc.core_properties
        metadata['author'] = core_props.author
        metadata['title'] = core_props.title
        metadata['subject'] = core_props.subject
        metadata['keywords'] = core_props.keywords
        metadata['created'] = core_props.created
        metadata['modified'] = core_props.modified
        metadata['last_modified_by'] = core_props.last_modified_by
        
        # Document structure
        metadata['paragraphs'] = len(doc.paragraphs)
        metadata['tables'] = len(doc.tables)
        
        return metadata
        
    except Exception as e:
        logger.error(f"Failed to extract DOCX metadata: {str(e)}")
        return metadata


def extract_paragraphs_with_style(docx_path: str) -> List[dict]:
    """
    Extract paragraphs with their style information
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        List of dictionaries with text and style information
    """
    if not DOCX_AVAILABLE:
        return []
    
    try:
        doc = Document(docx_path)
        paragraphs_with_style = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs_with_style.append({
                    'text': para.text,
                    'style': para.style.name if para.style else None,
                    'alignment': str(para.alignment) if para.alignment else None
                })
        
        return paragraphs_with_style
        
    except Exception as e:
        logger.error(f"Failed to extract paragraphs with style: {str(e)}")
        return []


def count_words(docx_path: str) -> int:
    """
    Count words in DOCX file
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        Word count
    """
    try:
        text = extract_text_from_docx(docx_path)
        words = text.split()
        return len(words)
    except Exception:
        return 0


def extract_headings(docx_path: str) -> List[dict]:
    """
    Extract all headings from DOCX file
    
    Args:
        docx_path: Path to DOCX file
        
    Returns:
        List of dictionaries with heading text and level
    """
    if not DOCX_AVAILABLE:
        return []
    
    try:
        doc = Document(docx_path)
        headings = []
        
        for para in doc.paragraphs:
            if para.style and para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.split()[-1])
                    headings.append({
                        'text': para.text,
                        'level': level
                    })
                except (ValueError, IndexError):
                    headings.append({
                        'text': para.text,
                        'level': 1
                    })
        
        return headings
        
    except Exception as e:
        logger.error(f"Failed to extract headings: {str(e)}")
        return []


# Alias for backwards compatibility
parse_docx = extract_text_from_docx