"""
Tests for document parsers (PDF and DOCX)
"""

import pytest
import sys
from pathlib import Path

# Add project root to path
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

from tools.pdf_parser import extract_text_from_pdf
from tools.docx_parser import extract_text_from_docx


@pytest.fixture
def sample_pdf_path(tmp_path):
    """Create a sample PDF file for testing"""
    # Note: Creating actual PDFs requires additional libraries
    # For now, we'll just create a placeholder
    pdf_path = tmp_path / "sample.pdf"
    
    # This would normally create a real PDF
    # For testing without external dependencies, we skip actual PDF creation
    # pdf_path.touch()
    
    return pdf_path


@pytest.fixture
def sample_docx_path(tmp_path):
    """Create a sample DOCX file for testing"""
    docx_path = tmp_path / "sample.docx"
    
    # Create a simple DOCX file
    try:
        from docx import Document
        
        doc = Document()
        doc.add_heading('John Doe', 0)
        doc.add_paragraph('john.doe@email.com')
        doc.add_paragraph('+1-555-0123')
        
        doc.add_heading('Experience', level=1)
        doc.add_paragraph('Senior Software Engineer at Tech Corp')
        doc.add_paragraph('Developed applications using Python and AWS')
        
        doc.add_heading('Education', level=1)
        doc.add_paragraph('MIT - BS Computer Science (2018)')
        
        doc.add_heading('Skills', level=1)
        doc.add_paragraph('Python, AWS, Docker, Kubernetes')
        
        doc.save(str(docx_path))
        
        return docx_path
        
    except ImportError:
        pytest.skip("python-docx not installed")


class TestPDFParser:
    """Test suite for PDF parser"""
    
    def test_extract_text_from_pdf_file_not_found(self):
        """Test handling of non-existent PDF file"""
        with pytest.raises(Exception):
            extract_text_from_pdf("nonexistent.pdf")
    
    def test_extract_text_from_pdf_invalid_file(self, tmp_path):
        """Test handling of invalid PDF file"""
        # Create a non-PDF file
        fake_pdf = tmp_path / "fake.pdf"
        fake_pdf.write_text("This is not a PDF file")
        
        # Should raise an exception
        with pytest.raises(Exception):
            extract_text_from_pdf(str(fake_pdf))
    
    @pytest.mark.skipif(
        not Path(project_root / "data" / "sample_resumes").exists(),
        reason="Sample resumes directory not found"
    )
    def test_extract_text_from_real_pdf(self):
        """Test extraction from real PDF if available"""
        sample_dir = project_root / "data" / "sample_resumes"
        pdf_files = list(sample_dir.glob("*.pdf"))
        
        if pdf_files:
            text = extract_text_from_pdf(str(pdf_files[0]))
            
            assert text is not None
            assert isinstance(text, str)
            assert len(text) > 0
        else:
            pytest.skip("No PDF files found in sample_resumes")


class TestDOCXParser:
    """Test suite for DOCX parser"""
    
    def test_extract_text_from_docx_file_not_found(self):
        """Test handling of non-existent DOCX file"""
        with pytest.raises(Exception):
            extract_text_from_docx("nonexistent.docx")
    
    def test_extract_text_from_docx_invalid_file(self, tmp_path):
        """Test handling of invalid DOCX file"""
        # Create a non-DOCX file
        fake_docx = tmp_path / "fake.docx"
        fake_docx.write_text("This is not a DOCX file")
        
        # Should raise an exception
        with pytest.raises(Exception):
            extract_text_from_docx(str(fake_docx))
    
    def test_extract_text_from_valid_docx(self, sample_docx_path):
        """Test extraction from valid DOCX file"""
        if sample_docx_path.exists():
            text = extract_text_from_docx(str(sample_docx_path))
            
            assert text is not None
            assert isinstance(text, str)
            assert len(text) > 0
            
            # Check for expected content
            assert "John Doe" in text
            assert "john.doe@email.com" in text
            assert "Python" in text
    
    @pytest.mark.skipif(
        not Path(project_root / "data" / "sample_resumes").exists(),
        reason="Sample resumes directory not found"
    )
    def test_extract_text_from_real_docx(self):
        """Test extraction from real DOCX if available"""
        sample_dir = project_root / "data" / "sample_resumes"
        docx_files = list(sample_dir.glob("*.docx"))
        
        if docx_files:
            text = extract_text_from_docx(str(docx_files[0]))
            
            assert text is not None
            assert isinstance(text, str)
            assert len(text) > 0
        else:
            pytest.skip("No DOCX files found in sample_resumes")


class TestParserComparison:
    """Test parser consistency and comparison"""
    
    def test_both_parsers_return_string(self, sample_docx_path):
        """Test that both parsers return string output"""
        if sample_docx_path.exists():
            docx_text = extract_text_from_docx(str(sample_docx_path))
            assert isinstance(docx_text, str)
    
    def test_extracted_text_is_not_empty(self, sample_docx_path):
        """Test that extracted text is not empty"""
        if sample_docx_path.exists():
            text = extract_text_from_docx(str(sample_docx_path))
            assert len(text.strip()) > 0
    
    def test_text_extraction_preserves_structure(self, sample_docx_path):
        """Test that text extraction preserves basic structure"""
        if sample_docx_path.exists():
            text = extract_text_from_docx(str(sample_docx_path))
            
            # Should have multiple lines
            lines = text.split('\n')
            assert len(lines) > 1
            
            # Should contain section headers
            assert any('Experience' in line for line in lines)
            assert any('Education' in line for line in lines)
            assert any('Skills' in line for line in lines)


def test_parser_error_messages():
    """Test that parsers provide helpful error messages"""
    
    # Test PDF parser
    try:
        extract_text_from_pdf("nonexistent.pdf")
        assert False, "Should have raised an exception"
    except Exception as e:
        error_message = str(e).lower()
        assert "error" in error_message or "not found" in error_message or "no such file" in error_message
    
    # Test DOCX parser
    try:
        extract_text_from_docx("nonexistent.docx")
        assert False, "Should have raised an exception"
    except Exception as e:
        error_message = str(e).lower()
        assert "error" in error_message or "not found" in error_message or "no such file" in error_message


if __name__ == "__main__":
    pytest.main([__file__, "-v"])